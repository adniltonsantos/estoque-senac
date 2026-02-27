import os
from docx import Document

# python files to include
files = [
    '/workspaces/estoque-senac/data.py',
    '/workspaces/estoque-senac/utils.py',
    '/workspaces/estoque-senac/main.py',
    '/workspaces/estoque-senac/produtos/__init__.py',
    '/workspaces/estoque-senac/produtos/cadastro.py',
    '/workspaces/estoque-senac/produtos/listagem.py',
    '/workspaces/estoque-senac/produtos/edicao.py',
    '/workspaces/estoque-senac/produtos/deletar.py',
    '/workspaces/estoque-senac/produtos/calculos.py',
]

# create document
doc = Document()
doc.add_heading('Código Fonte - estoque-senac', level=1)

for filepath in files:
    if not os.path.exists(filepath):
        continue
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            code = f.read()
    except Exception as e:
        print(f"Erro ao ler {filepath}: {e}")
        continue
    
    filename = filepath.split('/')[-1]
    doc.add_heading(filename, level=2)
    # add code block with monospace font
    para = doc.add_paragraph()
    para.style = 'Normal'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = 100000  # 8pt

output = '/workspaces/estoque-senac/source_code.docx'
doc.save(output)
print(f'Document gerado em: {output}')
